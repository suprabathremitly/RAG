"""Document processing service for handling uploads and text extraction."""
import os
import uuid
from pathlib import Path
from typing import List, Dict, Any, Tuple
from datetime import datetime
import logging

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

# Document parsers
import pypdf
from docx import Document as DocxDocument

from app.config import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles document upload, parsing, and chunking."""
    
    SUPPORTED_FORMATS = {'.pdf', '.txt', '.docx'}
    
    def __init__(self):
        self.upload_dir = Path(settings.upload_directory)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    def is_supported_format(self, filename: str) -> bool:
        """Check if file format is supported."""
        return Path(filename).suffix.lower() in self.SUPPORTED_FORMATS
    
    async def process_upload(self, file_content: bytes, filename: str) -> Tuple[str, List[Document]]:
        """
        Process uploaded file: save, extract text, and chunk.
        
        Returns:
            Tuple of (document_id, list of chunked documents)
        """
        if not self.is_supported_format(filename):
            raise ValueError(f"Unsupported file format. Supported: {self.SUPPORTED_FORMATS}")
        
        # Generate unique document ID
        doc_id = str(uuid.uuid4())
        
        # Save file
        file_path = self.upload_dir / f"{doc_id}_{filename}"
        with open(file_path, 'wb') as f:
            f.write(file_content)
        
        logger.info(f"Saved document {filename} with ID {doc_id}")
        
        # Extract text
        text = self._extract_text(file_path)
        
        # Create metadata
        metadata = {
            'document_id': doc_id,
            'filename': filename,
            'file_path': str(file_path),
            'file_size': len(file_content),
            'file_type': Path(filename).suffix.lower(),
            'upload_timestamp': datetime.utcnow().isoformat()
        }
        
        # Chunk the document
        chunks = self._chunk_document(text, metadata)
        
        logger.info(f"Created {len(chunks)} chunks from document {filename}")
        
        return doc_id, chunks
    
    def _extract_text(self, file_path: Path) -> str:
        """Extract text from document based on file type."""
        suffix = file_path.suffix.lower()
        
        try:
            if suffix == '.pdf':
                return self._extract_from_pdf(file_path)
            elif suffix == '.txt':
                return self._extract_from_txt(file_path)
            elif suffix == '.docx':
                return self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {suffix}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise
    
    def _extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        text = []
        with open(file_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text.strip():
                    text.append(f"[Page {page_num + 1}]\n{page_text}")
        
        return "\n\n".join(text)
    
    def _extract_from_txt(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read()
    
    def _extract_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        doc = DocxDocument(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        return "\n\n".join(paragraphs)
    
    def _chunk_document(self, text: str, metadata: Dict[str, Any]) -> List[Document]:
        """Split document into chunks with metadata."""
        # Split text into chunks
        text_chunks = self.text_splitter.split_text(text)
        
        # Create Document objects with metadata
        documents = []
        for i, chunk in enumerate(text_chunks):
            chunk_metadata = metadata.copy()
            chunk_metadata['chunk_id'] = f"{metadata['document_id']}_chunk_{i}"
            chunk_metadata['chunk_index'] = i
            chunk_metadata['total_chunks'] = len(text_chunks)
            
            documents.append(Document(
                page_content=chunk,
                metadata=chunk_metadata
            ))
        
        return documents
    
    def get_document_info(self, doc_id: str) -> Dict[str, Any]:
        """Get information about a stored document."""
        # Find file with this doc_id
        for file_path in self.upload_dir.glob(f"{doc_id}_*"):
            stat = file_path.stat()
            return {
                'document_id': doc_id,
                'filename': file_path.name.replace(f"{doc_id}_", ""),
                'file_size': stat.st_size,
                'file_type': file_path.suffix.lower(),
                'file_path': str(file_path),
                'upload_timestamp': datetime.fromtimestamp(stat.st_ctime)
            }
        
        return None
    
    def list_documents(self) -> List[Dict[str, Any]]:
        """List all stored documents."""
        documents = []
        seen_ids = set()
        
        for file_path in self.upload_dir.glob("*"):
            if file_path.is_file():
                # Extract doc_id from filename (format: {doc_id}_{original_filename})
                parts = file_path.name.split("_", 1)
                if len(parts) == 2:
                    doc_id = parts[0]
                    if doc_id not in seen_ids:
                        seen_ids.add(doc_id)
                        info = self.get_document_info(doc_id)
                        if info:
                            documents.append(info)
        
        return documents
    
    def delete_document(self, doc_id: str) -> bool:
        """Delete a document and its file."""
        for file_path in self.upload_dir.glob(f"{doc_id}_*"):
            try:
                file_path.unlink()
                logger.info(f"Deleted document file: {file_path}")
                return True
            except Exception as e:
                logger.error(f"Error deleting document {doc_id}: {e}")
                return False
        
        return False


# Global instance
document_processor = DocumentProcessor()

